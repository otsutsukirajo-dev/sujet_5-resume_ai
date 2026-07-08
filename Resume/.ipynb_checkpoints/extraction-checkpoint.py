

import os
from typing import Optional

import pdfplumber
from docx import Document as DocxDocument


class ExtractionError(Exception):
    """Levée quand l'extraction de texte échoue."""
    pass


ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt"}


def allowed_file(filename: str) -> bool:
    """Vérifie si l'extension du fichier est supportée."""
    ext = os.path.splitext(filename)[1].lower()
    return ext in ALLOWED_EXTENSIONS


def extract_text(filepath: str) -> str:
    """
    Point d'entrée principal : détecte le type de fichier et
    délègue à la bonne fonction d'extraction.

    Args:
        filepath: chemin absolu ou relatif vers le fichier.

    Returns:
        Le texte extrait (str), potentiellement vide si le fichier
        ne contient pas de texte détectable.

    Raises:
        ExtractionError: si le format n'est pas supporté ou si
        l'extraction échoue.
    """
    if not os.path.exists(filepath):
        raise ExtractionError(f"Fichier introuvable : {filepath}")

    ext = os.path.splitext(filepath)[1].lower()

    if ext == ".pdf":
        return _extract_from_pdf(filepath)
    elif ext == ".docx":
        return _extract_from_docx(filepath)
    elif ext == ".txt":
        return _extract_from_txt(filepath)
    else:
        raise ExtractionError(f"Extension non supportée : {ext}")


def _extract_from_pdf(filepath: str) -> str:
    """Extrait le texte d'un PDF page par page avec pdfplumber."""
    text_parts = []
    try:
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
    except Exception as e:
        raise ExtractionError(f"Erreur lors de la lecture du PDF : {e}")

    text = "\n".join(text_parts).strip()
    if not text:
        raise ExtractionError(
            "Aucun texte détecté dans ce PDF (peut-être un scan/image "
            "sans OCR)."
        )
    return text


def _extract_from_docx(filepath: str) -> str:
    """Extrait le texte d'un fichier Word (.docx), paragraphes + tableaux."""
    try:
        doc = DocxDocument(filepath)
    except Exception as e:
        raise ExtractionError(f"Erreur lors de la lecture du Word : {e}")

    parts = [p.text for p in doc.paragraphs if p.text.strip()]

    # On récupère aussi le texte dans les tableaux, souvent oublié
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())

    text = "\n".join(parts).strip()
    if not text:
        raise ExtractionError("Aucun texte détecté dans ce document Word.")
    return text


def _extract_from_txt(filepath: str, encoding: Optional[str] = None) -> str:
    """Lit un fichier texte brut, avec fallback d'encodage."""
    encodings_to_try = [encoding] if encoding else ["utf-8", "latin-1"]

    last_error = None
    for enc in encodings_to_try:
        try:
            with open(filepath, "r", encoding=enc) as f:
                text = f.read().strip()
            if not text:
                raise ExtractionError("Le fichier texte est vide.")
            return text
        except UnicodeDecodeError as e:
            last_error = e
            continue

    raise ExtractionError(f"Impossible de décoder le fichier texte : {last_error}")